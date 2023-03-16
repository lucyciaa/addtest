from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
import time
from docx import Document

def crawler_3():
    document = Document()
    options = Options()
    options.add_experimental_option("debuggerAddress", "127.0.0.1:9527")
    browser = webdriver.Chrome(options=options)

    url_list = []
    filename = './url.txt'
    file = open(filename, 'r', encoding='utf-8')
    for line in file.readlines():
        url_list.append('http' + line.split('http')[1].split('，')[0])
    for url in url_list:
        print(url)
        browser.get(url)
        time.sleep(0.5)
        # 输出：标题
        title = browser.find_elements(by=By.CLASS_NAME, value='title')
        print('标题：' + title[0].text)
        # 输出：内容
        desc = browser.find_elements(by=By.CLASS_NAME, value='desc')

        videoFlag = browser.find_elements(by=By.ID, value='videoPlayer')
        if(len(videoFlag) > 0):
            document.add_paragraph('视频')
        document.add_paragraph(desc[0].text)
        file_title = title[0].text
        if(file_title.__contains__('|')):
            file_title = file_title.replace('|','｜')
        document.save(file_title + '.docx')


if __name__ == '__main__':
    crawler_3()